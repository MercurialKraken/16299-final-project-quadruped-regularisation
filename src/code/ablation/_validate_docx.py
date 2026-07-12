"""Quick validity check on the deliverable docx."""
import zipfile

from docx import Document

PATH = r"A:\Ablation_Deliverables\Ablation_Report.docx"

# 1) zip integrity
with zipfile.ZipFile(PATH, "r") as zf:
    bad = zf.testzip()
    print(f"zip integrity: {'OK' if bad is None else f'BAD: {bad}'}")
    print(f"member count: {len(zf.namelist())}")
    has_doc = any("word/document.xml" in n for n in zf.namelist())
    has_image = any(n.startswith("word/media/") for n in zf.namelist())
    print(f"  word/document.xml: {has_doc}")
    print(f"  word/media/* (images): {has_image}")

# 2) python-docx parses cleanly
doc = Document(PATH)
print(f"OPENABLE: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
print(f"first 5 paragraphs:")
for i, p in enumerate(doc.paragraphs[:5]):
    print(f"  [{i}] {p.text[:80]!r}")
print(f"first table dims: {len(doc.tables[0].rows)}x{len(doc.tables[0].columns)}" if doc.tables else "no tables")
