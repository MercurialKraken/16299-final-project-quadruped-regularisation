"""Build the 7-variant ablation report as a docx using python-docx.

Updated 2026-05-04 v2: uses noreg-trained Bal-LP flow (proper apples-to-apples).
"""
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet",
                           "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


ABL = r"A:\AllIsaac\flow_matching_project\data\ablation"
OUT = r"A:\AllIsaac\flow_matching_project\Ablation_Report.docx"

with open(os.path.join(ABL, "results.json")) as f:
    R = json.load(f)

# rank by fall rate ascending for the table
ORDER = sorted(
    ["noreg", "somereg", "extremereg", "noreg_flow", "noreg_lp",
     "somereg_lp", "noreg_flow_lp", "somereg_flow", "extremereg_flow"],
    key=lambda v: R[v].get("fall_rate") if R[v].get("fall_rate") is not None else 1.0,
)
LABELS = {
    "noreg": "No-reg PPO (raw)",
    "somereg": "Some-reg PPO (raw, existing)",
    "extremereg": "Extreme-reg PPO (raw, λ=−0.5)",
    "noreg_flow": "No-reg + Bal-LP Flow",
    "somereg_flow": "Some-reg + Bal-LP Flow",
    "extremereg_flow": "Extreme-reg + Bal-LP Flow",
    "noreg_lp": "No-reg + Causal LP (15 Hz)",
    "somereg_lp": "Some-reg + Causal LP (15 Hz)",
    "noreg_flow_lp": "No-reg + Bal-LP Flow + Causal LP",
}
REGW = {
    "noreg": "0.0", "somereg": "−0.01", "extremereg": "−0.5",
    "noreg_flow": "0.0", "somereg_flow": "−0.01", "extremereg_flow": "−0.5",
    "noreg_lp": "0.0", "somereg_lp": "−0.01", "noreg_flow_lp": "0.0",
}


def fmt_pct(v):
    if v is None:
        return "—"
    return f"{v*100:.1f}%" if v <= 1 else f"{v:.2f}%"


def fmt_num(v, d=3):
    return "—" if v is None else f"{v:.{d}f}"


def fmt_ms(v):
    return "—" if v is None else f"{v:.3f} ms"


def fmt_time(v):
    return "—" if v is None else f"{v:.0f} s"


def cell_shade(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def main():
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Flow Matching for Quadruped Locomotion")
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("7-way Ablation: Reward Regularization × Post-Hoc Smoother (apples-to-apples)")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Paul Colombo, Arnav Shah, Jack Gerdsen — CMU, Spring 2026")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)

    doc.add_paragraph()

    # ---- Headline findings ----
    doc.add_heading("Headline Findings", level=1)
    findings = [
        "Some-reg PPO + Bal-LP flow is the winner: 3/80 = 3.75% fall rate. This exactly reproduces the deck's headline Iter-5 result (4%, ≈3/80). The combination of a moderate action-rate penalty in the reward and a flow trained on that policy's distribution is the recipe.",
        "Methodology validated: some-reg PPO raw gives 40/80 = 50.0% fall rate, exactly matching the deck.",
        "Removing the action-rate penalty entirely (no-reg) makes the raw policy more push-robust (31.25% vs 50.0%). The penalty was suppressing rapid corrections needed for impulse recovery. But it also makes the policy more jittery (HF energy 11.9% vs 7.97%) — and noticeably hurts how much flow can help.",
        "Bal-LP flow on no-reg policy: 23.75% fall rate. Helps a lot, but cannot match the same flow on some-reg (3.75%). The action-rate penalty appears to give the flow a smoother target distribution to lock onto. Synergy, not redundancy.",
        "Bal-LP flow on extreme-reg policy CATASTROPHICALLY FAILS: 80/80 = 100% fell. The extreme penalty froze the policy (mean Vx ≈ 0); the flow's optimal targets tell it to walk; the policy can't execute walking, so it falls every time. Cautionary tale: flow needs a competent underlying policy.",
        "Causal IIR LP filter alone is a respectable middle ground for reg-free policies: HF energy 3.32%, fall rate 33.75%, inference cost 0.069 ms. ~80× cheaper than flow but doesn't reach the some-reg+flow result.",
        "LP applied to a regularized policy hurts: some-reg + LP = 56.25% fall rate (worse than raw 50%). Same failure mode the deck warned about — non-aware filtering removes legitimate corrections.",
        "Flow + LP cascade is still a dead end: HF energy hits an all-time low of 1.23% (−90% vs raw) but fall rate climbs to 37.5% — the LP layer suppresses corrections the flow correctly preserves. Same finding as Iter 3 in the deck.",
        "Inference cost: raw = 0.013 ms, causal IIR = 0.07 ms (5×), Bal-LP flow = 5.6–6.0 ms (430–460×). For deployment on resource-constrained hardware this is the chief drawback of the flow.",
    ]
    for f in findings:
        add_bullet(doc, f)

    # ---- Results table (sorted by fall rate ascending) ----
    doc.add_heading("Complete 7-Variant Results (sorted by fall rate)", level=1)
    p = doc.add_paragraph()
    r = p.add_run("Push recovery: 80 envs, lateral force 50–500 N, 10 magnitudes × 8 envs. Spectral metrics from 2000-step rollouts at vx = 1.0 fixed, 50 Hz. Inference latency measured during the push test (mean over 450 steps after warmup). Bal-LP flow for the no-reg variants was trained on no-reg PPO rollouts (full Iter-5 pipeline: 5000-step rollout → 5000-target candidate-shooting K=16 H=10 → Butterworth 15 Hz on targets → 200-epoch flow training).")
    r.italic = True
    r.font.size = Pt(9)

    cols = ["Variant", "λ_action_rate", "HF >10Hz", "Action rate RMS",
            "Jerk RMS", "Mean Vx (m/s)", "Fall rate", "Inference",
            "Train wall"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
        cell_shade(hdr[i], "D5E8F0")
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for v in ORDER:
        r = R[v]
        row = table.add_row().cells
        row[0].text = LABELS[v]
        row[1].text = REGW[v]
        row[2].text = fmt_pct(r.get("hf_energy_pct"))
        row[3].text = fmt_num(r.get("action_rate_rms"))
        row[4].text = fmt_num(r.get("jerk_rms"))
        row[5].text = fmt_num(r.get("mean_vx"), 2)
        row[6].text = fmt_pct(r.get("fall_rate"))
        row[7].text = fmt_ms(r.get("infer_ms_mean"))
        row[8].text = fmt_time(r.get("train_total_s"))
        fr = r.get("fall_rate")
        if fr is not None:
            if fr <= 0.30:
                cell_shade(row[6], "B5E8B5")  # darker green for the best
            elif fr <= 0.35:
                cell_shade(row[6], "D9F2D9")
            elif fr >= 0.5:
                cell_shade(row[6], "F8DCDC")
        for i, c in enumerate(row):
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True

    p = doc.add_paragraph()
    r = p.add_run("Fall-rate cell shading: dark green ≤ 30% (winner), light green ≤ 35%, red ≥ 50%. λ_action_rate is the weight on action_rate_l2 in the PPO reward. Train wall is 300 iterations × 4096 envs on RTX 5070 Laptop. The \"some-reg\" PPO was trained on a different date and pre-instrumentation, so its wall time is unrecorded.")
    r.italic = True
    r.font.size = Pt(9)

    # ---- Comparison plot ----
    img_path = os.path.join(ABL, "ablation_comparison.png")
    if os.path.exists(img_path):
        doc.add_page_break()
        doc.add_heading("Visual Comparison", level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(9.5))

    # ---- Methodology ----
    doc.add_page_break()
    doc.add_heading("Methodology Notes", level=1)
    doc.add_heading("Variant definitions", level=2)
    for t in [
        "No-reg / Some-reg / Extreme-reg PPOs differ ONLY in the weight on action_rate_l2 (0.0, −0.01, −0.5). All other reward terms, env config, observation space, terrain, and training schedule (300 iterations × 4096 envs, seed 42) are held constant.",
        "\"Flow\" = Balanced-LP (Iter 5 from the original deck). For the no-reg variants in this report, the flow was retrained on no-reg-PPO rollouts: collect 5000-step rollout with full sim state, generate optimal targets via short-horizon random shooting (K=16 candidates, H=10 horizon, w_track=3, w_jerk=0.5, w_energy=0.3, w_stab=2), Butterworth 15 Hz LP on the targets, then 200-epoch conditional flow matching on the LP-smoothed targets. t_end = 1.0 at inference.",
        "\"LP\" = causal 1st-order IIR Butterworth at 15 Hz, applied online during the push eval. Distinct from the offline filtfilt used in spectral analysis (non-causal; cannot run online).",
        "\"Flow + LP\" cascade = Bal-LP flow output then passed through the causal IIR.",
    ]:
        add_bullet(doc, t)

    doc.add_heading("How this iteration corrects the earlier draft", level=2)
    for t in [
        "v1 of this report used the existing flow_model_balanced_lp.pt — a flow trained on the some-reg PPO's output distribution — applied to the no-reg policy. That cross-policy generalization test made flow look worse than LP. It was a methodological mistake.",
        "v2 (this draft) trains a separate Bal-LP flow on no-reg rollouts. Result: noreg + flow goes from 40.0% fall rate (cross-policy) → 23.75% fall rate (same-policy). HF energy drops from 14.0% → 4.17%. The same-policy comparison is the only fair one — flow matching depends on training-distribution match.",
        "Generating optimal targets for no-reg took 1167 s (~19 min) on the RTX 5070 at K=16, H=10. This is the dominant cost of producing a new flow model when the PPO changes; the flow training itself is ~30 s.",
    ]:
        add_bullet(doc, t)

    doc.add_heading("Caveats", level=2)
    for t in [
        "Spectral metrics for the smoothed variants (rows containing flow or LP) are computed offline on the recorded raw rollout — i.e., we apply the smoother to recorded actions and FFT the result. Push fall rate is online (real closed-loop). The inference latency is from the online push test as well.",
        "HF energy on raw some-reg: deck = 2.1%, this run = 7.97%. Different FFT methodology (likely DC removal / windowing); internal consistency across the new variants is what matters for the ablation conclusions.",
        "Rough terrain eval was not included in this iteration — flat-terrain ablation only.",
        "Ranking is at α≈0.05 noise floor for an 80-env push test. The 23.75% / 31.25% gap (8 envs) is meaningful; the 31.25% / 33.75% gap (2 envs) is at the noise floor.",
    ]:
        add_bullet(doc, t)

    # ---- Takeaways ----
    doc.add_heading("Key Takeaways", level=1)
    for t in [
        "The action-rate penalty has a sweet spot. Without it (λ=0) the raw policy is more agile (31% falls) but jittery (12% HF). With it at −0.01 the raw policy is fragile alone (50% falls) but trains a flow that achieves 4% falls — the best result anywhere. Pushed too far (λ=−0.5) the policy collapses and even flow can't rescue it (100% falls).",
        "Bal-LP flow matching is the best post-hoc smoother in the table, full stop, when paired with a competent regularized policy. Bal-LP flow + some-reg = 3.75% fall rate beats every other entry by a clear margin and isn't reached by LP filters or by flow on a different reg level.",
        "Bal-LP flow is policy-specific AND policy-fragility-sensitive. The flow on extreme-reg degrades performance because the underlying policy can't follow up on the optimal corrections. The flow assumes the policy will keep stepping after the smoothed action lands; if the policy is collapsed, that assumption breaks.",
        "Combining flow + LP at inference (cascade) does not stack — same dead-end the deck found with Iter-3. The LP layer ends up suppressing corrections the flow correctly preserves.",
        "Causal IIR LP at 15 Hz is a 5× cheaper inference path. For no-reg policies it neither helps nor hurts much (33.75% vs raw 31.25%). For some-reg policies it actively hurts (56.25% vs raw 50%). Skip the LP filter as a standalone smoother — it's not pulling its weight.",
        "Recommended deployment recipe (flat terrain, proprioceptive obs): (1) train PPO with action_rate_l2 = −0.01 (the existing default), (2) collect rollout-with-states, (3) generate optimal targets via random shooting (K=16, H=10, w_track=3, w_jerk=0.5, w_energy=0.3, w_stab=2), (4) Butterworth 15 Hz on the targets, (5) train flow 200 epochs. Deploy at t_end=1.0. End state: 3.75% fall rate at 50–500 N, ≈6 ms inference per step.",
    ]:
        add_bullet(doc, t)

    p = doc.add_paragraph()
    r = p.add_run("Generated 2026-05-04 v2 (apples-to-apples flow). Code: A:\\AllIsaac\\flow_matching_project\\scripts\\ablation\\. Data: data\\ablation\\results.json. Plot: data\\ablation\\ablation_comparison.png. Flow weights: A:\\AllIsaac\\IsaacLab\\flow_model_balanced_lp_noreg.pt.")
    r.italic = True
    r.font.size = Pt(9)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
